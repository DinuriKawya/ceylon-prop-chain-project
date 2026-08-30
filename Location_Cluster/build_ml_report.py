from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(r"C:\Users\VULCAN\Desktop\2026\Location_Cluster")
OUTPUT = BASE / "ML_Location_Analysis_Implementation.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "5B6573"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_run_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True

    if "Figure Analysis" not in [s.name for s in doc.styles]:
        analysis = doc.styles.add_style("Figure Analysis", WD_STYLE_TYPE.PARAGRAPH)
    else:
        analysis = doc.styles["Figure Analysis"]
    analysis.font.name = "Calibri"
    analysis._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    analysis._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    analysis.font.size = Pt(9.5)
    analysis.font.color.rgb = RGBColor.from_string(INK)
    analysis.paragraph_format.space_after = Pt(5)
    analysis.paragraph_format.line_spacing = 1.05

    # Quiet running header and page number footer for a long report section.
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("CEYLONPROPCHAIN  |  MACHINE LEARNING LOCATION ANALYSIS")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)
    p_pr = hp._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D7DBE2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Page ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_page_number(fp)


def add_title_page(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(92)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(14)
    set_run_font(kicker.add_run("IMPLEMENTATION REPORT"), size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    title.paragraph_format.keep_with_next = True
    set_run_font(title.add_run("Machine Learning\nLocation Analysis Module"), size=27, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(54)
    set_run_font(subtitle.add_run("CeylonPropChain"), size=15, color=MUTED)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Inches(0.55)
    summary.paragraph_format.right_indent = Inches(0.55)
    summary.paragraph_format.space_after = Pt(72)
    set_run_font(
        summary.add_run(
            "This module converts historical Sri Lankan apartment advertisements into precomputed city- and district-level location profiles. "
            "The implementation uses unsupervised K-means clustering, price-per-square-foot statistics, geographic enrichment, and JSON export for the React application."
        ),
        size=11,
        color=INK,
    )

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(date.add_run("Final implementation record"), size=10, color=MUTED, italic=True)
    doc.add_page_break()


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2)
    else:
        set_run_font(p.add_run(text))
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run_font(p.add_run(text))
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(header)), size=9.5, bold=True)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(str(value)), size=9.2)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_figure(doc, filename, caption, purpose, observation, decision, width=6.05):
    path = BASE / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(path), width=Inches(width))

    c = doc.add_paragraph(style="Caption")
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.add_run(caption)

    a = doc.add_paragraph(style="Figure Analysis")
    for label, value in (("Purpose: ", purpose), ("Observation: ", observation), ("Decision: ", decision)):
        lead = a.add_run(label)
        set_run_font(lead, size=9.5, bold=True)
        body = a.add_run(value)
        set_run_font(body, size=9.5)
    return a


def add_new_page(doc):
    doc.add_page_break()


def build():
    doc = Document()
    style_document(doc)
    add_title_page(doc)

    add_heading(doc, "1. Module purpose and implementation approach", 1)
    add_para(
        doc,
        "The module provides a reproducible location-analysis layer for the property platform. It does not predict an individual apartment price. Instead, it groups cities with similar historical average prices per square foot, enriches those groups with listing activity, year-on-year change, postal codes and coordinates, and exports the results for application use."
    )
    add_para(
        doc,
        "Unsupervised learning was appropriate because the source data did not contain a verified target label such as premium, mid-market or budget. K-means therefore discovered groups from the numerical price profile, after which descriptive names were assigned for display."
    )

    add_heading(doc, "1.1 End-to-end workflow", 2)
    add_table(
        doc,
        ["Stage", "Implementation"],
        [
            ("1. Data acquisition", "Load property advertisements and Sri Lankan city/district reference data from Kaggle."),
            ("2. Exploration", "Inspect listing types, property categories, locations and date coverage before selecting the modelling subset."),
            ("3. Cleaning", "Retain apartments for sale, parse price and floor area, derive price per square foot, and remove invalid values and outliers."),
            ("4. Aggregation", "Normalize city names and calculate city-level mean price per square foot and listing count; retain cities with at least three listings."),
            ("5. Clustering", "Standardize city average price per square foot and fit K-means with three clusters."),
            ("6. Enrichment", "Add district summaries, 2022-2023 trend, postal code, coordinates and similar-city suggestions."),
            ("7. Validation and export", "Test lookup paths, save the model/scaler and export four JSON files consumed by the React frontend."),
        ],
        [1800, 7560],
    )

    add_heading(doc, "1.2 Datasets", 2)
    add_table(
        doc,
        ["Dataset", "Local data used", "Role"],
        [
            ("Sri Lanka Property Ads Dataset", "203,874 records; 26 fields", "Historical advertisement data used to derive apartment price and activity statistics."),
            ("Sri Lanka Provinces, Districts and Cities", "2,155 city records and 25 district records", "Reference data used for district mapping, postal codes, latitude and longitude."),
        ],
        [2700, 1900, 4760],
    )
    src = doc.add_paragraph()
    src.paragraph_format.space_before = Pt(4)
    src.paragraph_format.space_after = Pt(4)
    set_run_font(src.add_run("Sources: "), size=9.5, color=MUTED, bold=True)
    set_run_font(src.add_run("https://www.kaggle.com/datasets/ivantha/sri-lanka-property-ads-dataset; "), size=9.5, color=MUTED)
    set_run_font(src.add_run("https://www.kaggle.com/datasets/tharindumadhusanka9/sri-lanka-provinces-districts-cities"), size=9.5, color=MUTED)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    set_run_font(note.add_run("Record counts refer to the local downloaded files used in this implementation."), size=9.5, color=MUTED, italic=True)

    add_new_page(doc)
    add_heading(doc, "1.3 Libraries", 2)
    add_table(
        doc,
        ["Library", "Use in the module"],
        [
            ("pandas", "CSV loading, filtering, cleaning, grouping, joining and JSON-ready tables."),
            ("NumPy", "Numerical operations and missing-value handling."),
            ("Matplotlib", "Exploratory, model-selection and result visualizations."),
            ("scikit-learn", "StandardScaler, KMeans, silhouette score and Davies-Bouldin index."),
            ("Joblib", "Saving and reloading the fitted scaler and clustering model."),
            ("json / os", "Structured output export and file handling."),
        ],
        [2100, 7260],
    )

    add_heading(doc, "2. Exploratory data analysis", 1)
    add_para(doc, "Exploration was performed before modelling to verify that the source data contained a usable apartment-for-sale subset and adequate location and date coverage.")
    add_figure(
        doc,
        "eda_type_split.png",
        "Figure 1. Distribution of listing transaction types in the raw property dataset.",
        "Confirm whether sale advertisements were sufficiently represented for a sale-price analysis. ",
        "For-sale records formed 168,971 of 203,874 listings (82.9%); for-rent records formed 33,376 (16.4%), while the remaining transaction labels were small. ",
        "Restrict the modelling dataset to sale listings so sale prices are not mixed with monthly rents.",
        width=5.2
    )

    add_new_page(doc)
    add_heading(doc, "2.1 Property-category coverage", 2)
    add_figure(
        doc,
        "eda_category_breakdown.png",
        "Figure 2. Most frequent property categories in the raw dataset.",
        "Determine whether apartments can be isolated from land, houses, rentals and commercial properties. ",
        "The data is dominated by land and house advertisements. Apartments for sale appear as a smaller but usable category (5,198 raw category records). ",
        "Filter by both transaction type and apartment category before numerical cleaning; this produced 5,177 candidate apartment-for-sale records."
    )

    add_new_page(doc)
    add_heading(doc, "2.2 Location coverage", 2)
    add_figure(
        doc,
        "eda_top_locations.png",
        "Figure 3. Highest-volume locations across all raw property types.",
        "Assess the geographic concentration and spelling structure of the location field. ",
        "Listing volume is uneven: Piliyandala (17,791) and Homagama (9,100) lead the raw data, followed by Padukka (7,472). ",
        "Normalize city text, aggregate at city level, and use a minimum-listing rule so very sparse cities do not create unstable profiles."
    )

    add_new_page(doc)
    add_heading(doc, "2.3 Time coverage", 2)
    add_figure(
        doc,
        "eda_year_coverage.png",
        "Figure 4. Listing volume by year in the raw dataset.",
        "Verify whether the data supports a comparable year-on-year indicator. ",
        "The dataset covers September 2022 to April 2023, with 79,936 dated records in 2022 and 123,932 in 2023. The periods are partial rather than full calendar years. ",
        "Calculate a descriptive 2022-to-2023 change only where a city has at least three usable records in each year, and report missing trend values rather than estimating them."
    )

    add_new_page(doc)
    add_heading(doc, "3. Data preprocessing and feature engineering", 1)
    add_para(doc, "The apartment-for-sale subset was cleaned in a fixed sequence so the price feature was comparable across records:")
    add_bullet(doc, "Currency symbols, commas and other text were removed from the price field; values at or below LKR 1,000,000 were excluded as implausible total apartment sale prices.")
    add_bullet(doc, "Floor area was extracted from the semi-structured property-details field. Missing areas and values at or below 100 square feet were removed.")
    add_bullet(doc, "The modelling feature was calculated as price_per_sqft = cleaned sale price / extracted floor area. No square-root transformation was applied.")
    add_bullet(doc, "The 1.5 x IQR rule removed extreme price-per-square-foot values. The usable row count decreased from 5,152 to 4,693.")
    add_bullet(doc, "City names were standardized, city averages and listing counts were calculated, and only cities with at least three usable listings were retained. This produced 44 city profiles.")
    add_figure(
        doc,
        "chart_outlier_removal.png",
        "Figure 5. Price-per-square-foot distribution before and after IQR outlier removal.",
        "Show how extreme values affected the modelling feature and verify the effect of the cleaning rule. ",
        "The raw distribution contained a long upper tail; the IQR-filtered distribution retained the main market range while removing extreme values. ",
        "Fit the city summaries and K-means model using the 4,693 cleaned apartment records."
    )

    add_new_page(doc)
    add_heading(doc, "4. Unsupervised learning and K-means clustering", 1)
    add_para(doc, "K-means was applied to one modelling feature: each city's average apartment price per square foot. StandardScaler transformed that feature to a common standardized scale before training. K-means was selected because the implementation required understandable, non-overlapping city groups and the dataset did not provide verified class labels.")
    add_heading(doc, "4.1 Selecting the number of clusters", 2)
    add_figure(
        doc,
        "chart_elbow_method.png",
        "Figure 6. Elbow analysis for candidate values of k.",
        "Compare within-cluster sum of squares as the number of clusters increases. ",
        "Inertia decreases for every added cluster, with the largest practical reduction occurring in the early values and a visible bend around three groups. ",
        "Use k = 3 as a compact segmentation that remains straightforward to interpret in the property interface."
    )

    add_new_page(doc)
    add_heading(doc, "4.2 Internal validation", 2)
    add_figure(
        doc,
        "chart_silhouette.png",
        "Figure 7. Silhouette profile for the selected three-cluster solution.",
        "Inspect the cohesion and separation of individual cities after fitting k = 3. ",
        "The average silhouette coefficient was 0.511. Most city values were positive, although some were near zero and therefore lay close to a neighbouring cluster boundary. ",
        "Accept the three-cluster solution as interpretable but acknowledge that its separation is moderate rather than perfect."
    )
    add_para(doc, "Across the candidate values, k = 2 produced the highest silhouette score (0.592). The selected k = 3 score was 0.511. Three groups were retained because the elbow evidence and the usefulness of three reportable price tiers outweighed the modest reduction in silhouette score; k = 3 is therefore not presented as mathematically optimal.")
    add_table(
        doc,
        ["k", "Silhouette score", "Davies-Bouldin index"],
        [
            ("2", "0.592", "0.593"),
            ("3 (selected)", "0.511", "0.570"),
            ("4", "0.517", "0.537"),
            ("5", "0.534", "0.479"),
            ("6", "0.547", "0.457"),
        ],
        [1800, 3300, 4260],
    )

    add_new_page(doc)
    add_heading(doc, "5. Cluster creation and interpretation", 1)
    add_para(doc, "The fitted cluster identifiers were not assumed to have a meaningful order. The learned centroids were sorted from lowest to highest price, and the presentation labels were then assigned consistently to the low-, middle- and high-price groups.")
    add_table(
        doc,
        ["Display label", "Cities", "Mean LKR/sq ft", "Observed range (LKR/sq ft)"],
        [
            ("Budget & Rising", "21", "21,111.07", "7,598.34 - 26,715.47"),
            ("Fast Growing", "17", "32,462.24", "27,597.85 - 40,037.42"),
            ("Expensive & Stable", "6", "52,198.76", "44,146.84 - 62,377.57"),
        ],
        [2400, 1200, 2200, 3560],
    )
    label_note = doc.add_paragraph()
    label_note.paragraph_format.space_before = Pt(4)
    label_note.paragraph_format.space_after = Pt(6)
    set_run_font(label_note.add_run("Important interpretation: "), size=9.5, color=DARK_BLUE, bold=True)
    set_run_font(label_note.add_run("the model used price only. The words 'Rising', 'Growing' and 'Stable' are interface descriptions added after clustering; growth and stability were not features learned by K-means."), size=9.5, color=INK)
    add_figure(
        doc,
        "chart_price_by_city.png",
        "Figure 8. Average apartment price per square foot by city and assigned cluster.",
        "Display the learned separation and the relative price position of every retained city. ",
        "The cities form three ordered price bands with six cities in the highest band, seventeen in the middle band and twenty-one in the lowest band. ",
        "Use the ordered centroids, rather than raw K-means ID values, to assign stable display labels."
    )

    add_new_page(doc)
    add_heading(doc, "5.1 Price and listing activity", 2)
    add_figure(
        doc,
        "chart_cluster_scatter.png",
        "Figure 9. City price clusters plotted against listing activity.",
        "Inspect how the price-based groups relate to the amount of supporting data available for each city. ",
        "Listing counts vary substantially within the same price group, showing that activity was not a K-means input and should not be interpreted as a clustering boundary. ",
        "Retain listing_count as a separate confidence and popularity indicator in the exported location profile."
    )

    add_new_page(doc)
    add_heading(doc, "5.2 Cluster balance", 2)
    add_figure(
        doc,
        "chart_cluster_distribution.png",
        "Figure 10. Distribution of the 44 retained cities across the three clusters.",
        "Check whether the chosen segmentation collapses most cities into a single class. ",
        "The result is uneven but usable: 21 cities are in the budget group, 17 in the middle group and 6 in the expensive group. ",
        "Keep the learned class sizes and report them transparently rather than forcing equal-sized groups."
    )

    add_new_page(doc)
    add_heading(doc, "6. Enrichment of the cluster output", 1)
    add_para(doc, "The price clusters were extended with information required by the application. District profiles were classified with the same fitted scaler and model; the model was not retrained for districts. City records were joined to postal codes and coordinates. Forty-one cities matched the reference data automatically, while Angoda, Kotte and Kiribathgoda were completed with manual reference values. Similar-city suggestions were selected from the same cluster and ranked by listing count.")
    add_heading(doc, "6.1 District summaries", 2)
    add_figure(
        doc,
        "chart_district_prices.png",
        "Figure 11. District-level average apartment price per square foot.",
        "Provide a broader fallback when a requested city is unavailable. ",
        "Six districts met the data requirement. Colombo had both the largest evidence base (4,503 listings) and the highest district average; several other districts were represented by much smaller samples. ",
        "Return district-level results only as a fallback and include listing count so the evidence level remains visible."
    )

    add_new_page(doc)
    add_heading(doc, "6.2 Geographic coverage", 2)
    add_figure(
        doc,
        "chart_geographic_map.png",
        "Figure 12. Geographic distribution of the enriched city profiles, including a Colombo-area inset.",
        "Verify coordinate enrichment and show the spatial concentration of the final city set. ",
        "All 44 city profiles contain coordinates, with strong concentration around the Colombo metropolitan area. The inset separates the most active nearby cities that would overlap on the national view. ",
        "Export coordinates for map display, while verifying the three manually supplied records against an authoritative source before production deployment."
    )

    add_new_page(doc)
    add_heading(doc, "6.3 Trend completeness", 2)
    add_figure(
        doc,
        "chart_trend_completeness.png",
        "Figure 13. Availability of the 2022-to-2023 descriptive trend indicator.",
        "Quantify how much of the city output supports a year-on-year comparison after the minimum sample rule. ",
        "Twenty-eight of 44 cities (64%) had at least three usable listings in both years; sixteen cities did not have a defensible trend value. ",
        "Export missing trends as unavailable rather than filling or extrapolating them. Treat the trend as descriptive because both years cover partial periods."
    )

    add_new_page(doc)
    add_heading(doc, "7. Testing, export and application integration", 1)
    add_heading(doc, "7.1 Lookup testing", 2)
    add_para(doc, "The final notebook reloaded the saved scaler and K-means model and tested the complete lookup function through five paths: exact city name, partial city name, postal code, district fallback, and unknown or empty input. Assertions verified that the output schema was complete and that invalid inputs returned a controlled response instead of an exception.")

    add_heading(doc, "7.2 Generated artefacts", 2)
    add_table(
        doc,
        ["Artefact", "Contents / purpose"],
        [
            ("kmeans_location_model.pkl", "Fitted three-cluster K-means model."),
            ("location_scaler.pkl", "StandardScaler fitted to city average price per square foot."),
            ("city_clusters.json", "44 application-ready city profiles with cluster, price, activity, trend, location and similar-city fields."),
            ("city_map_data.json", "44 compact city records for map visualization."),
            ("district_clusters.json", "Six district fallback profiles."),
            ("postal_to_city.json", "Postal-code lookup for the 44 final cities."),
        ],
        [2700, 6660],
    )

    add_heading(doc, "7.3 React integration", 2)
    add_para(doc, "The frontend consumes the precomputed city_clusters.json file through its ML service and returns a location profile during property and location workflows. The Python model is not executed in the browser, is not exposed as a live prediction API, and is not stored or executed on the blockchain. This separation keeps the blockchain module responsible for ownership and transaction logic while the ML module supplies read-only location intelligence.")

    add_heading(doc, "8. Limitations", 1)
    add_bullet(doc, "The clustering uses one feature—city average apartment price per square foot—so the groups represent price tiers rather than a complete measure of investment quality.")
    add_bullet(doc, "Advertisement prices are asking prices, not verified completed-sale values.")
    add_bullet(doc, "The 2022 and 2023 samples cover partial calendar periods; the trend value is descriptive and unavailable for 16 cities.")
    add_bullet(doc, "District samples outside Colombo are small, and three manually completed city coordinates/postal records require authoritative production verification.")
    add_bullet(doc, "The saved model is suitable for reproducible batch regeneration, while the current application intentionally uses the exported JSON rather than live inference.")

    add_heading(doc, "9. Implementation outcome", 1)
    add_para(doc, "The completed pipeline transformed 203,874 raw property advertisements into 4,693 cleaned apartment records, 44 city profiles, six district profiles and four frontend JSON files. Its main analytical result is a transparent three-tier city segmentation based on standardized average price per square foot, supported by listing activity, limited year-on-year trend information, geographic enrichment and tested lookup behaviour.")

    doc.core_properties.title = "Machine Learning Location Analysis Module"
    doc.core_properties.subject = "CeylonPropChain implementation report"
    doc.core_properties.author = "CeylonPropChain Project"
    doc.core_properties.keywords = "machine learning, K-means, location analysis, property, Sri Lanka"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
